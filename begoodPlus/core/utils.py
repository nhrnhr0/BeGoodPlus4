from begoodPlus.secrects import GOOGLE_SERVICE_ACCOUNT_FILE, SECRECT_CLIENT_SIDE_DOMAIN, ALL_MORDER_FILE_SPREEDSHEET_URL
from ordered_model.models import OrderedModelBase
import gspread
from oauth2client.service_account import ServiceAccountCredentials
from base64 import urlsafe_b64decode, urlsafe_b64encode
from uuid import UUID
import docx
from django.shortcuts import redirect
from docx.oxml.ns import qn
import os
import re
from docx.shared import Inches, Cm
import google_auth_oauthlib
from begoodPlus.secrects import GOOGLE_CLIENT_SECRET_PATH
from productSize.models import ProductSize
from django.conf import settings
from docx import Document
import datetime
from docx.shared import Pt
from docx.enum.table import WD_TABLE_DIRECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches
import googleapiclient
from googleapiclient.http import MediaFileUpload, MediaIoBaseDownload
import pandas as pd
import io
from docx.oxml import OxmlElement
from django.urls import reverse
from googleapiclient.discovery import build

from catalogAlbum.models import CatalogAlbum, TopLevelCategory
from catalogImages.models import CatalogImage
import gspread
from oauth2client.service_account import ServiceAccountCredentials
from begoodPlus.secrects import GOOGLE_SERVICE_ACCOUNT_FILE


def build_drive_service(cred):
    service = build('drive', 'v3', credentials=cred)
    return service


def url_to_edit_object(obj):
    url = reverse('admin:%s_%s_change' %
                  (obj._meta.app_label,  obj._meta.model_name),  args=[obj.id])
    return url


def fixUniqeSlug(apps=None, schema_editor=None):
    all_albums = CatalogAlbum.objects.all()
    for album in all_albums:
        album.save()

    all_top_levels = TopLevelCategory.objects.all()
    for top_level in all_top_levels:
        top_level.save()

    all_catalog_images = CatalogImage.objects.all()
    for catalog_image in all_catalog_images:
        catalog_image.save()


def get_drive_file(service, file_id):
    # pylint: disable=maybe-no-member
    request = service.files().export(fileId=file_id,
                                     mimeType='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    file = io.BytesIO()
    downloader = MediaIoBaseDownload(file, request)
    done = False
    try:
        while done is False:
            status, done = downloader.next_chunk()
            print(F'Download {int(status.progress() * 100)}.')
        return file.getvalue()
    except Exception as e:
        return str(e)


def get_sheet_from_drive_url(url, drive_service, drive_creds=None, loaded_files={}):

    fileId = url.split(
        'https://docs.google.com/spreadsheets/d/')[1].split('/edit')[0]
    sheetId = url.split('gid=')[1]
    if loaded_files.get(fileId):
        all_sheets = loaded_files.get(fileId)
    else:
        bytes_exel_file = get_drive_file(drive_service, fileId)
        # conver bytes to in memory file
        # if it's a string (error) return it
        if isinstance(bytes_exel_file, str):
            return bytes_exel_file, None, None
        file = io.BytesIO(bytes_exel_file)
        # process the file
        all_sheets = pd.ExcelFile(file)
        loaded_files[fileId] = all_sheets
    # f'https://docs.google.com/spreadsheets/d/{doc_id}/gviz/tq?tqx=out:csv&sheet={sheet_name}'
    # get sheetname from url
    sheetname = get_sheetname_from_driveurl(url, drive_creds)
    sheetname = sheetname[:31]
    # print(all_sheets.sheet_names)
    return all_sheets.parse(sheetname, header=0, dtype=str), sheetname, loaded_files


def get_sheetname_from_driveurl(url, drive_creds=None):
    sheetId = url.split('gid=')[1]
    http_client = googleapiclient.discovery._auth.authorized_http(
        drive_creds)
    response, content = http_client.request(url)
    cont = content.decode('utf-8')
    regex = sheetId
    regex += r"\\\",\[{\\\"\d\\\":\[\[\d,\d,\\\"(.+?)\\\"]"
    matches = re.search(regex, cont)
    sheet_name = matches.group(1)
    return sheet_name.replace('\\', '')


def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def add_table_to_doc(document, data):
    # , style="ColorfulList-Accent5"
    # First row are table headers!
    # https://github.com/python-openxml/python-docx/issues/149
    table = document.add_table(
        rows=(data.shape[0]+1), cols=data.shape[1], style="Light Shading")
    table.direction = WD_TABLE_DIRECTION.RTL
    table.autofit = True
    table.allow_autofit = True

    # widths = (Inches(3), Inches(3),)
    # for row in table.rows:
    #     for idx, width in enumerate(widths):
    #         row.cells[idx].width = width
    # table.direction = WD_TABLE_DIRECTION.LTR
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # first row is the headers (column names)
    for j in range(data.shape[-1]):
        table.cell(
            0, data.shape[-1] - j - 1).text = data.columns[j]
        table.cell(
            0, data.shape[-1] - j - 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # data rows
    for i in range(data.shape[0]):
        for j in range(data.shape[-1]):
            txt = str(data.values[data.shape[0] -
                                  i - 1, data.shape[-1] - j - 1])
            if txt == 'nan' or txt == '0' or txt == 'ON COLOR' or txt == '0.0':
                txt = ''
            table.cell(i + 1,
                       j).text = txt.replace('.0', '')
            table.cell(i + 1,
                       j).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell = table.cell(i + 1, j)
            set_cell_border(
                cell,
                top={"sz": 12, "color": "#000000", "val": "single"},
                bottom={"sz": 12, "color": "#000000", "val": "single"},
                left={"sz": 12, "color": "#000000", "val": "single"},
                right={"sz": 12, "color": "#000000", "val": "single"},
            )
            # if it's the last 3 columns, then align right
            has_modal = 'מודל' in data.columns.tolist()
            headers_len = data.shape[-1] - \
                3 if has_modal else data.shape[-1] - 2

            if j >= headers_len:
                table.cell(i + 1,
                           j).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # for i, column in enumerate(data):
    #     for row in range(data.shape[0]):
    #         table.cell(row, i).text = str(data[column][row])
    # for cell in table.columns[0].cells:
    #     cell.width = Inches(0.5)
    table.columns[len(table.columns) - 1].width = Inches(1.5)
    if (len(table.columns) > 2):
        table.columns[len(table.columns) - 2].width = Inches(0.95)
        table.columns[len(table.columns) - 3].width = Inches(0.95)
    # align center
    table.columns[len(table.columns) - 1].alignment = WD_TABLE_ALIGNMENT.RIGHT
    # table.rows[0].cells[0].width = Inches(5.0)
    return table


def generate_provider_docx(provider_data, provider_name):
    pd.option_context('expand_frame_repr', False, 'display.max_rows', None)
    document = Document()

    today = datetime.datetime.now()
    date_time = today.strftime("%d/%m/%Y")
    end_file_location = 'static_cdn' if settings.DEBUG else 'static'
    end_file_location = end_file_location + \
        '/assets/images/provider_docx_header.png'
    file_location = os.path.join(
        settings.BASE_DIR, end_file_location)

    document.add_picture(file_location, width=Cm(21.5 - 0.75*2))
    entries = []
    _last_product_name = ''
    for product_name, product_data in provider_data['products'].items():
        _last_product_name = product_name
        print(product_name)
        for item in product_data['items']:
            print('qty', item['qty'])
            if item['qty'] != '0' and item['qty'] != 0.0:
                entries.append({
                    'מוצר': product_name,
                    'מידה': str('ONE SIZE' if item['size'] == 'one size' else item['size']),
                    'צבע': str(item['color']),
                    'מודל': str(item['verient']),
                    'כמות': item['qty'],
                })

    # document = Document()
    # crete pivot table
    indexs = ['מוצר', 'מודל', 'צבע']
    column = 'מידה'
    value = 'כמות'
    df = pd.DataFrame(entries)
    df = df.fillna(0)
    # if df is empty, then return
    if df.empty:
        return {}, False
    df = df.pivot_table(index=indexs, columns=[column],
                        values=value, aggfunc='sum')
    df = df.fillna(0)
    df = df.astype(int)
    df = df.reset_index()

    options = ProductSize.objects.all().order_by(
        'code').values_list('size', flat=True)
    opt1 = ['XS', 'S',
            'M', 'L', 'XL', '2XL', '3XL', '4XL', '5XL']
    opt2 = ['36', '37', '38', '39',
            '40', '41', '42', '43', '44', '45', '46', '47', '48']
    opt3 = ['ONE SIZE']
    # opt4 = 2-18 (in tows)
    opt4 = ['2', '4', '6', '8', '10', '12', '14', '16', '18']
    # opt5 = 36-54 (in tows)
    opt5 = ['36', '38', '40', '42', '44', '46', '48', '50', '52', '54']
    # opt3 = ['מוצר', 'צבע', 'מודל', + all options exept what is in opt1 and opt2]
    opt6 = []
    for option in options:
        if option not in opt1 and option not in opt2 and option not in opt3 and option not in opt4 and option not in opt5:
            opt6.append(option)
    all_options = [opt1, opt2, opt3, opt4, opt5, opt6]
    # df = find for each product (מוצר) the best option to be used as the columns names (מידה)
    # for example if the product has only XS and S then the option will be opt1
    # if the product has only 36 and 37 then the option will be opt2
    # if the product has only ONE SIZE then the option will be opt3
    # if the product has only 2 and 4 then the option will be opt4
    # if the product has only 36 and 38 then the option will be opt2
    # if the product has only 36 and 38 and 50 then the option will be opt5
    # if the product has only 36 and 38 and 50 and 52 then the option will be opt5
    # code:
    tables = []
    options_dfs = {}
    for index, row in df.iterrows():
        # print(row)
        product_name = row['מוצר']
        # itate over all options and find the best one
        best_option = None
        best_option_len = 0
        best_option_idx = -1
        curr_option_idx = 0
        for option in all_options:
            found = 0

            indexs = row.index.to_list()
            indexs = list(filter(lambda x: x not in [
                'מוצר', 'מודל', 'צבע'], indexs))
            for size in option:
                if size in indexs and row.get(size) != 0:
                    found += 1
            if found > best_option_len:
                best_option_len = found
                best_option = option
                best_option_idx = curr_option_idx
            curr_option_idx += 1
        # create new df with the best option
        if options_dfs.get(best_option_idx) is None:
            options_dfs[best_option_idx] = pd.DataFrame()
        print(product_name, best_option)
        if best_option:
            d = {
                'מוצר': product_name,
                'מודל': row['מודל'],
                'צבע': row['צבע'],
                **{size: row.get('ONE SIZE' if size == 'one size' else size, '') for size in best_option}
            }
        else:
            d = {
                'מוצר': product_name,
                'מודל': row['מודל'],
                'צבע': row['צבע'],
                'ONE SIZE': row.get('ONE SIZE', '0')
            }
        options_dfs[best_option_idx] = options_dfs[best_option_idx].append(
            d, ignore_index=True)
    # base = ['מוצר', 'צבע', 'מודל']

    # opt1 = base + opt1[:: -1]
    # opt2 = base + opt2[:: -1]
    # t1 = df.reindex(labels=opt1, axis=1)
    # t2 = df.reindex(labels=opt2, axis=1)

    # #
    # t3 = df.reindex(labels=['מוצר', 'צבע', 'מודל', 'ONE SIZE'], axis=1)
    # add to t3 all what is in df that is not in t1 and t2
    styles_element = document.styles.element
    rpr_default = styles_element.xpath('./w:docDefaults/w:rPrDefault/w:rPr')[0]
    lang_default = rpr_default.xpath('w:lang')[0]
    lang_default.set(docx.oxml.shared.qn('w:val'), 'HE-IL')

    rtlstyle = document.styles.add_style('rtl', WD_STYLE_TYPE.PARAGRAPH)
    rtlstyle.font.rtl = True
    p = document.add_heading(
        'לכבוד: ' + provider_name + ' \t\t תאריך: ' + date_time, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # t1.dropna(axis=1, how='all', inplace=True)
    # if not t1.dropna(axis=1, how='all').empty:
    #     document.add_heading('טבלת ביגוד', level=2)
    #     add_table_to_doc(document, t1.dropna(axis=0, how='all'))
    # if not t2.dropna(axis=1, how='all').empty:
    #     document.add_heading('טבלת נעליים', level=2)
    #     add_table_to_doc(document, t2.dropna(axis=0, how='all'))
    # if not t3.dropna(axis=1, how='all').empty:
    #     document.add_heading('טבלת כלים', level=2)
    #     add_table_to_doc(document, t3.dropna(axis=0, how='all'))
    for key, value in options_dfs.items():
        if not value.dropna(axis=1, how='all').empty:
            cols = all_options[key][0]
            if key == 0:
                label = 'טבלת ביגוד'
            elif key == 1:
                label = 'טבלת נעליים'
            elif key == 2:
                label = 'טבלת כלים'
            elif key == 3:
                label = 'טבלת ילדים'
            elif key == 4:
                label = 'טבלת מכנסיים'
            else:
                label = 'טבלה'
            cols = value.columns.to_list()
            # remove 'מוצר', 'מודל', 'צבע' from cols
            cols = list(filter(lambda x: x not in [
                'מוצר', 'מודל', 'צבע'], cols))
            # iterate over all cols and remove all rows that have 0 in all cols from the start and the end but not in the middle
            cols.sort(
                key=lambda x: ProductSize.objects.get(size=x).code)
            to_remove = set()
            for col in cols:
                if value[col].sum() == 0 or value[col].sum() == '':
                    to_remove.add(col)
                else:
                    break
            for col in reversed(cols):
                if value[col].sum() == 0 or value[col].sum() == '':
                    to_remove.add(col)
                else:
                    break
            for col in to_remove:
                print('removing', col, ' from ',
                      label, value.columns.to_list())
                value.drop(col, axis=1, inplace=True)
                cols.remove(col)

            # add title with the label in the center
            p = document.add_heading(label, level=2)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            base = ['מוצר', 'צבע', 'מודל']
            # if 'ONE SIZE' in value.columns:
            #     base = ['מוצר', ]
            lbls = base + cols[::-1]  # all_options[key][:: -1]
            value = value.reindex(labels=lbls, axis=1)
            # if all the col of מודל are empty strings then remove the col
            models = pd.Series(value.get('מודל'), dtype='str').str.strip()
            found_model = False
            for m in models:
                if m != '' and m != 'nan' and m != '0' and m != 'None' and m != '0.0':
                    found_model = True
                    break
            if not found_model:
                try:
                    value = value.drop('מודל', axis=1)
                except:
                    pass

            add_table_to_doc(document, value.dropna(axis=0, how='all'))
    # changing the page margins
    margin = 1
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(margin)
        section.bottom_margin = Cm(margin)
        section.left_margin = Cm(margin)
        section.right_margin = Cm(margin)
    return document, True


def merge_data_to_providers_dict(original_data, provider_name, product_name, color, size, varient, qyt, morder_id):
    # data: {
    # '$provider_name': {
    # 'morders': [$morder_id1, $morder_id2, ...],
    # 'products': {
    # '$product_name': {
    # items = [
    # {
    # 'color': '$product_color',
    # 'size': '$product_size',
    # 'varient': '$product_varient',
    # 'amount': '$product_amount',
    # },
    # ]
    # }
    # }
    # }
    if provider_name not in original_data:
        original_data[provider_name] = {
            'products': {},
            'morders': [],
        }
    if morder_id not in original_data[provider_name]['morders']:
        original_data[provider_name]['morders'].append(morder_id)
    if product_name not in original_data[provider_name]['products']:
        original_data[provider_name]['products'][product_name] = {
            'items': [],
        }
    found = False
    for item in original_data[provider_name]['products'][product_name]['items']:
        if item['color'] == color and item['size'] == size and item['verient'] == varient:
            item['qty'] += qyt
            found = True
            break
    if not found:
        original_data[provider_name]['products'][product_name]['items'].append({
            'color': color,
            'size': size,
            'verient': varient,
            'qty': qyt,
        })
    return original_data


def process_sheets_to_providers_docx(sheets, obj):
    # beutifly print each sheet to console

    all_sheets_data = []
    sheet_index = 1
    for sheet in sheets:
        values = sheet.get_all_values()
        morders_id = str(int(float(values[1][0])))
        obj.logs.append('processing sheet: ' +
                        str(sheet_index) + ' morder: ' + morders_id)

        sheet_index += 1
        rows_data = []
        current_row_data = None
        sheet_data = values[3:]
        for idx, row in enumerate(sheet_data):

            # print(idx, ') ')
            # for i in range(len(sheet.columns)):
            #     print(row[i], end=' ')
            # if idx == 0 or idx == 1:
            #     # print('skip')
            #     continue
            # col[7] is the 'print?' column
            # if it has any value, it means we are on a header row
            if not pd.isna(row[7]) and row[7] != '':
                # print('header row')
                if current_row_data:
                    rows_data.append(current_row_data)
                product_name = row[1]
                obj.logs.append(
                    'row: ' + str(idx) + ' processing product: ' + product_name)
                header_total_amount = int(float(row[2]))
                header_taken_amount = row[4]
                if str(header_taken_amount).lower() == 'v':
                    header_taken_amount = header_total_amount
                elif pd.isna(header_taken_amount) or header_taken_amount == '':
                    header_taken_amount = 0
                else:
                    header_taken_amount = int(float(header_taken_amount))

                header_provider_name = row[11] if len(row) > 11 else ''
                header_amount = header_total_amount - header_taken_amount
                current_row_data = {
                    'product_name': product_name,
                    'header_anount': header_amount,
                    'header_provider_name': header_provider_name,
                    'has_child': False,
                    'morder_id': morders_id,
                }
                continue
            else:
                print('product row')
                product_color = row[0]
                product_size = row[1]
                product_varient = row[2]
                product_total_amount = int(
                    float(row[3] if row[3] != '' else 0))
                product_taken_amount = row[4]
                if str(product_taken_amount).lower() == 'v':
                    product_taken_amount = int(float(row[3]))
                elif pd.isna(product_taken_amount) or product_taken_amount == '':
                    product_taken_amount = 0
                else:
                    product_taken_amount = int(float(product_taken_amount))
                product_provider_name = row[11] if len(row) > 11 else ''
                obj.logs.append('row: ' + str(idx) + ' processing product: ' + str(product_name) + ' color: ' + str(product_color) + ' size: ' + str(product_size) +
                                ' varient: ' + str(product_varient) + ' total_amount: ' + str(product_total_amount) + ' taken_amount: ' + str(product_taken_amount))
                print(product_taken_amount, product_total_amount)
                print(type(product_taken_amount), type(product_total_amount))
                current_row_data['has_child'] = True
                if(current_row_data.get('childs') == None):
                    current_row_data['childs'] = []
                current_row_data['childs'].append({
                    'product_color': product_color,
                    'product_size': product_size,
                    'product_varient': product_varient,
                    'product_amount': product_total_amount - product_taken_amount,
                    'product_provider_name': product_provider_name,
                })
                continue
        # adding the last row
        if current_row_data:
            rows_data.append(current_row_data)

        for row in rows_data:
            # filter out the rows that has no childs and the header amount is 0 or less and rows childs that has 0 or less amount
            all_sheets_data.append(row)
        obj.logs.append('finished processing sheet: ' + str(sheet_index))
    # merge data as needed
    data = {}
    print(all_sheets_data)
    obj.logs.append('merging data')
    # obj.logs.append(all_sheets_data)
    for row in all_sheets_data:
        if not row['has_child']:
            provider_name = row['header_provider_name'] if not pd.isna(
                row['header_provider_name']) else ''
            size = 'ONE SIZE'
            color = 'NO COLOR'
            varient = ''
            qyt = row['header_anount']
            product_name = row['product_name']
            morder_id = row['morder_id']
            if qyt <= 0:
                continue
            data = merge_data_to_providers_dict(
                data, provider_name, product_name, color, size, varient, qyt, morder_id)
        else:
            product_name = row['product_name']
            morder_id = row['morder_id']
            for child in row['childs']:
                provider_name = child['product_provider_name'] if not pd.isna(
                    child['product_provider_name']) else ''
                size = child['product_size']
                color = child['product_color']
                varient = child['product_varient']
                qyt = child['product_amount']
                if qyt <= 0:
                    continue
                data = merge_data_to_providers_dict(
                    data, provider_name, product_name, color, size, varient, qyt, morder_id)
    obj.logs.append('finished merging data')
    # obj.logs.append(data)
    return data


def uuid2slug(uuidstring):
    if uuidstring:
        if isinstance(uuidstring, str):
            try:
                return urlsafe_b64encode(bytearray.fromhex(uuidstring)).rstrip(b'=').decode('ascii')
            except:
                return urlsafe_b64encode(str.encode(uuidstring)).rstrip(b'=').decode('ascii')
        else:
            return urlsafe_b64encode(uuidstring.bytes).rstrip(b'=').decode('ascii')
    else:
        return '<error>'


def slug2uuid(slug):
    return str(UUID(bytes=urlsafe_b64decode(slug + '==')))


def number_to_spreedsheet_letter(number):
    number = number - 1
    letter = ''
    while number >= 0:
        letter = chr((number % 26) + ord('A')) + letter
        number = number // 26 - 1
    return letter
